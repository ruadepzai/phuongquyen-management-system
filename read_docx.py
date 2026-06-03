import docx
import sys

doc = docx.Document(r"D:\WS học\TTCN\Đề cương Thực tập chuyên ngành - Nhóm 14.docx")

with open(r"C:\Users\Admin\docs_clean.md", "w", encoding="utf-8") as f:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            f.write(text + "\n")
    
    # Also extract tables
    for i, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {i+1} ---\n")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

print("Done")
